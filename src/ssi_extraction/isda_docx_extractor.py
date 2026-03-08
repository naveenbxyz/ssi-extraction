from __future__ import annotations

import logging
from pathlib import Path
import re
from typing import Optional, Tuple, Union

logger = logging.getLogger(__name__)


def _clean(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").strip().split())


def _normalize_attribute_key(text: str) -> str:
    key = _clean(text)
    # Common table formatting has a trailing ":" after attribute names.
    key = re.sub(r"[:\uFF1A]+\s*$", "", key).strip()
    return key


def _is_serial_number(text: str) -> bool:
    token = text.strip()
    if not token:
        return False
    return re.fullmatch(r"\d+[.)]?", token) is not None


def _parse_serial_number(text: str) -> Optional[int]:
    token = text.strip()
    if not token:
        return None
    match = re.match(r"^\s*(\d+)\s*[.)]?\s*$", token)
    if not match:
        return None
    try:
        return int(match.group(1))
    except ValueError:
        return None


def _extract_question_number(serial_cell: str, key_cell: str) -> Optional[int]:
    serial_num = _parse_serial_number(serial_cell)
    if serial_num is not None:
        return serial_num

    key_match = re.match(r"^\s*(\d+)\s*[.)-]\s*", key_cell or "")
    if key_match:
        try:
            return int(key_match.group(1))
        except ValueError:
            return None
    return None


def _strip_question_number_prefix(key: str) -> str:
    return re.sub(r"^\s*\d+\s*[.)-]\s*", "", key).strip()


def _row_to_key_value(cells: list[str]) -> Optional[Tuple[str, str, Optional[int]]]:
    if not any(cells):
        return None

    # ISDA table convention (primary): col1 = serial number, col2 = attribute/question, col3+ = value.
    if len(cells) >= 3 and _is_serial_number(cells[0]):
        key = _strip_question_number_prefix(_normalize_attribute_key(cells[1]))
        value = " | ".join([c.strip() for c in cells[2:] if c.strip()]).strip()
        if value:
            return key, value, _extract_question_number(cells[0], cells[1])

    # Some documents may have blank first cell with key/value shifted right.
    if len(cells) >= 3 and not cells[0].strip() and cells[1].strip():
        key = _strip_question_number_prefix(_normalize_attribute_key(cells[1]))
        value = " | ".join([c.strip() for c in cells[2:] if c.strip()]).strip()
        if value:
            return key, value, _extract_question_number(cells[0], cells[1])

    populated = [c for c in cells if c]
    if len(populated) >= 2 and populated[0]:
        key = _strip_question_number_prefix(_normalize_attribute_key(populated[0]))
        value = " | ".join(populated[1:]).strip()
        if value:
            question_number = _extract_question_number(cells[0] if cells else "", populated[0])
            return key, value, question_number

    if len(populated) == 1 and ":" in populated[0]:
        left, right = populated[0].split(":", 1)
        key = _clean(left)
        value = _clean(right)
        if key and value:
            question_number = _extract_question_number(cells[0] if cells else "", left)
            return key, value, question_number

    return None


def extract_docx_payload(docx_path: Union[str, Path]) -> dict:
    logger.info("ISDA DOCX extraction started path=%s", docx_path)
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise RuntimeError(
            "python-docx is required for ISDA extraction. Install with: pip install python-docx"
        ) from exc

    document = Document(str(docx_path))

    paragraphs: list[dict] = []
    for idx, para in enumerate(document.paragraphs, start=1):
        text = _clean(para.text)
        if text:
            paragraphs.append({"paragraph_index": idx, "text": text})

    tables: list[dict] = []
    key_value_candidates: list[dict] = []
    for t_idx, table in enumerate(document.tables, start=1):
        row_payloads: list[dict] = []
        for r_idx, row in enumerate(table.rows, start=1):
            cells = [_clean(cell.text) for cell in row.cells]
            if not any(cells):
                continue
            row_payloads.append({"row_index": r_idx, "cells": cells})
            kv = _row_to_key_value(cells)
            if kv:
                key_value_candidates.append(
                    {
                        "table_index": t_idx,
                        "row_index": r_idx,
                        "question_number": kv[2],
                        "key": kv[0],
                        "value": kv[1],
                    }
                )

        if row_payloads:
            tables.append({"table_index": t_idx, "rows": row_payloads})

    payload = {
        "paragraphs": paragraphs,
        "tables": tables,
        "key_value_candidates": key_value_candidates,
        "stats": {
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "candidate_count": len(key_value_candidates),
        },
    }
    logger.info(
        "ISDA DOCX extraction completed paragraphs=%d tables=%d kv_candidates=%d",
        len(paragraphs),
        len(tables),
        len(key_value_candidates),
    )
    return payload
